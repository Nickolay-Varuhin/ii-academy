"""Сервис генерации отчётов о достижениях пользователя.
Форматы: PDF (reportlab, кириллица через TTF) и DOCX (python-docx).
Стиль: чёрно-белый, формат ГОСТ 19 / ГОСТ Р 7.0.97.
"""
from __future__ import annotations

import io
import sys
import os
from datetime import datetime
from typing import Optional

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable,
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from docx import Document
from docx.shared import Pt, Cm as DCm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ════════════════════════════════════════════════════════════════════════
# Поиск TTF-шрифтов с кириллицей — перебираем все известные места
# Windows : Times New Roman или Arial в C:\Windows\Fonts\
# Linux   : Liberation Serif или DejaVu в /usr/share/fonts/truetype/
# ════════════════════════════════════════════════════════════════════════
_FONTS_REGISTERED = False


def _all_font_candidates() -> list:
    candidates = []
    win_fonts = os.path.join(os.environ.get("WINDIR", "C:\\Windows"), "Fonts")
    local_fonts = os.path.join(
        os.environ.get("LOCALAPPDATA", ""), "Microsoft", "Windows", "Fonts"
    )
    for fd in [win_fonts, local_fonts]:
        candidates.append({
            "regular":    os.path.join(fd, "times.ttf"),
            "bold":       os.path.join(fd, "timesbd.ttf"),
            "italic":     os.path.join(fd, "timesi.ttf"),
            "bolditalic": os.path.join(fd, "timesbi.ttf"),
        })
        candidates.append({
            "regular":    os.path.join(fd, "arial.ttf"),
            "bold":       os.path.join(fd, "arialbd.ttf"),
            "italic":     os.path.join(fd, "ariali.ttf"),
            "bolditalic": os.path.join(fd, "arialbi.ttf"),
        })
    for fd in [
        "/usr/share/fonts/truetype/liberation",
        "/usr/share/fonts/truetype/dejavu",
        "/usr/share/fonts/truetype/freefont",
    ]:
        bases = {
            "/usr/share/fonts/truetype/liberation": (
                "LiberationSerif-Regular.ttf", "LiberationSerif-Bold.ttf",
                "LiberationSerif-Italic.ttf", "LiberationSerif-BoldItalic.ttf",
            ),
            "/usr/share/fonts/truetype/dejavu": (
                "DejaVuSerif.ttf", "DejaVuSerif-Bold.ttf",
                "DejaVuSerif-Italic.ttf", "DejaVuSerif-BoldOblique.ttf",
            ),
            "/usr/share/fonts/truetype/freefont": (
                "FreeSerif.ttf", "FreeSerifBold.ttf",
                "FreeSerifItalic.ttf", "FreeSerifBoldItalic.ttf",
            ),
        }.get(fd)
        if bases:
            candidates.append({
                "regular":    os.path.join(fd, bases[0]),
                "bold":       os.path.join(fd, bases[1]),
                "italic":     os.path.join(fd, bases[2]),
                "bolditalic": os.path.join(fd, bases[3]),
            })
    return candidates


def _find_font_paths() -> dict:
    for c in _all_font_candidates():
        if all(os.path.isfile(v) for v in c.values()):
            return c
    tried = "\n".join(c["regular"] for c in _all_font_candidates())
    raise FileNotFoundError(
        "Не найдены TTF-шрифты с кириллицей.\nПроверены:\n" + tried
    )


def _ensure_fonts():
    global _FONTS_REGISTERED
    if _FONTS_REGISTERED:
        return
    paths = _find_font_paths()
    pdfmetrics.registerFont(TTFont("LiberSerif",    paths["regular"]))
    pdfmetrics.registerFont(TTFont("LiberSerif-B",  paths["bold"]))
    pdfmetrics.registerFont(TTFont("LiberSerif-I",  paths["italic"]))
    pdfmetrics.registerFont(TTFont("LiberSerif-BI", paths["bolditalic"]))
    from reportlab.pdfbase.pdfmetrics import registerFontFamily
    registerFontFamily(
        "LiberSerif",
        normal="LiberSerif", bold="LiberSerif-B",
        italic="LiberSerif-I", boldItalic="LiberSerif-BI",
    )
    _FONTS_REGISTERED = True


# ════════════════════════════════════════════════════════════════════════
# Data model
# ════════════════════════════════════════════════════════════════════════
class UserReportData:
    def __init__(
        self,
        user_id: int,
        last_name: str,
        first_name: str,
        patronymic: Optional[str],
        email: str,
        department: Optional[str],
        position: Optional[str],
        role: str,
        created_at: datetime,
        assignments_total: int,
        assignments_completed: int,
        assignments_in_progress: int,
        assignments_overdue: int,
        dialogs_total: int,
        dialogs_completed: int,
        avg_dialog_score: Optional[float],
        skill_levels: list,
        generated_at: datetime,
        generated_by: str,
    ):
        self.user_id              = user_id
        self.last_name            = last_name
        self.first_name           = first_name
        self.patronymic           = patronymic
        self.email                = email
        self.department           = department or "—"
        self.position             = position or "—"
        self.role                 = role
        self.created_at           = created_at
        self.assignments_total       = assignments_total
        self.assignments_completed   = assignments_completed
        self.assignments_in_progress = assignments_in_progress
        self.assignments_overdue     = assignments_overdue
        self.dialogs_total        = dialogs_total
        self.dialogs_completed    = dialogs_completed
        self.avg_dialog_score     = avg_dialog_score
        self.skill_levels         = skill_levels
        self.generated_at         = generated_at
        self.generated_by         = generated_by

    @property
    def full_name(self) -> str:
        parts = [self.last_name, self.first_name]
        if self.patronymic:
            parts.append(self.patronymic)
        return " ".join(parts)

    @property
    def short_name(self) -> str:
        parts = [self.last_name]
        if self.first_name:
            parts.append(self.first_name[0] + ".")
        if self.patronymic:
            parts.append(self.patronymic[0] + ".")
        return " ".join(parts)

    @property
    def completion_rate(self) -> float:
        if not self.assignments_total:
            return 0.0
        return round(self.assignments_completed / self.assignments_total * 100, 1)


# ════════════════════════════════════════════════════════════════════════
# PDF
# ════════════════════════════════════════════════════════════════════════
_BLACK = colors.black
_WHITE = colors.white
_LGRAY = colors.HexColor("#e8e8e8")
_MGRAY = colors.HexColor("#cccccc")


def _mastery_label(status: str) -> str:
    return {
        "mastered":    "Освоен",
        "in_progress": "В процессе",
        "not_started": "Не начат",
    }.get(status, status)


def _styles() -> dict:
    F = "LiberSerif"
    return {
        "doc_title": ParagraphStyle("DocTitle", fontName=F+"-B", fontSize=14,
            textColor=_BLACK, spaceAfter=4, alignment=TA_CENTER, leading=20),
        "org": ParagraphStyle("Org", fontName=F, fontSize=11,
            textColor=_BLACK, spaceAfter=4, alignment=TA_CENTER, leading=16),
        "section": ParagraphStyle("Section", fontName=F+"-B", fontSize=12,
            textColor=_BLACK, spaceBefore=16, spaceAfter=6, alignment=TA_LEFT, leading=16),
        "body": ParagraphStyle("Body", fontName=F, fontSize=12,
            textColor=_BLACK, spaceAfter=4, alignment=TA_JUSTIFY, leading=18,
            firstLineIndent=1.25*cm),
        "caption": ParagraphStyle("Caption", fontName=F, fontSize=11,
            textColor=_BLACK, spaceAfter=3, spaceBefore=10, alignment=TA_LEFT, leading=16),
        "center": ParagraphStyle("Center", fontName=F, fontSize=12,
            textColor=_BLACK, alignment=TA_CENTER, leading=18),
        "right": ParagraphStyle("Right", fontName=F, fontSize=11,
            textColor=_BLACK, alignment=TA_RIGHT),
        "tbl_hdr": ParagraphStyle("TblHdr", fontName=F+"-B", fontSize=11,
            textColor=_BLACK, alignment=TA_CENTER, leading=15),
        "tbl_cell": ParagraphStyle("TblCell", fontName=F, fontSize=11,
            textColor=_BLACK, alignment=TA_LEFT, leading=15),
        "tbl_cell_c": ParagraphStyle("TblCellC", fontName=F, fontSize=11,
            textColor=_BLACK, alignment=TA_CENTER, leading=15),
        "footer": ParagraphStyle("Footer", fontName=F+"-I", fontSize=10,
            textColor=_BLACK, alignment=TA_CENTER),
        "sub": ParagraphStyle("Sub", fontName=F+"-I", fontSize=12,
            alignment=TA_CENTER, spaceAfter=8, textColor=_BLACK),
    }


def _tbl_base_style() -> list:
    return [
        ("FONTNAME",      (0, 0), (-1, -1), "LiberSerif"),
        ("FONTSIZE",      (0, 0), (-1, -1), 11),
        ("FONTNAME",      (0, 0), (-1, 0),  "LiberSerif-B"),
        ("BACKGROUND",    (0, 0), (-1, 0),  _LGRAY),
        ("GRID",          (0, 0), (-1, -1), 0.8, _BLACK),
        ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING",    (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING",   (0, 0), (-1, -1), 6),
        ("RIGHTPADDING",  (0, 0), (-1, -1), 6),
    ]


def generate_pdf(data: UserReportData) -> bytes:
    _ensure_fonts()
    buf = io.BytesIO()
    L = R = 2.5 * cm
    T = B = 2.0 * cm
    W_page, _ = A4
    CW = W_page - L - R

    doc = SimpleDocTemplate(buf, pagesize=A4,
        leftMargin=L, rightMargin=R, topMargin=T, bottomMargin=B,
        title=f"Отчёт о достижениях — {data.full_name}")
    S = _styles()
    e = []

    # Шапка
    e.append(Paragraph(
        "Федеральное "
        "государственное "
        "бюджетное образо"
        "вательное учрежд"
        "ение<br/>высшего "
        "образования<br/>"
        "«Новгородский "
        "государственный "
        "университет "
        "имени Ярослава "
        "Мудрого»", S["org"]))
    e.append(Spacer(1, 4))
    e.append(HRFlowable(width="100%", thickness=1.5, color=_BLACK, spaceAfter=6))
    e.append(HRFlowable(width="100%", thickness=0.5, color=_BLACK, spaceAfter=10))
    e.append(Paragraph("ОТЧЁТ О ДОСТИЖЕНИЯХ", S["doc_title"]))
    e.append(Paragraph(
        "при прохождении "
        "курса в системе "
        "«ИИ-Академия»", S["sub"]))
    e.append(Spacer(1, 6))

    meta_rows = [
        [Paragraph("Составил:", S["tbl_cell"]),
         Paragraph(data.generated_by, S["tbl_cell"])],
        [Paragraph("Дата формирования:", S["tbl_cell"]),
         Paragraph(data.generated_at.strftime("%d.%m.%Y"), S["tbl_cell"])],
    ]
    meta_t = Table(meta_rows, colWidths=[5*cm, CW-5*cm], hAlign="RIGHT")
    meta_t.setStyle(TableStyle([
        ("FONTNAME",       (0,0),(-1,-1),"LiberSerif"),
        ("FONTSIZE",       (0,0),(-1,-1),11),
        ("FONTNAME",       (0,0),(0,-1), "LiberSerif-B"),
        ("GRID",           (0,0),(-1,-1),0.5,_MGRAY),
        ("TOPPADDING",     (0,0),(-1,-1),4),
        ("BOTTOMPADDING",  (0,0),(-1,-1),4),
        ("LEFTPADDING",    (0,0),(-1,-1),6),
    ]))
    e.append(meta_t)
    e.append(Spacer(1, 14))
    e.append(HRFlowable(width="100%", thickness=1, color=_BLACK, spaceAfter=10))

    # 1. Сведения о сотруднике
    e.append(Paragraph("1. Сведения о сотруднике", S["section"]))
    profile_data = [
        [Paragraph("Показатель", S["tbl_hdr"]),
         Paragraph("Значение", S["tbl_hdr"])],
        [Paragraph("Фамилия, имя, отчество", S["tbl_cell"]),
         Paragraph(data.full_name, S["tbl_cell"])],
        [Paragraph("Электронная почта", S["tbl_cell"]),
         Paragraph(data.email, S["tbl_cell"])],
        [Paragraph("Подразделение", S["tbl_cell"]),
         Paragraph(data.department, S["tbl_cell"])],
        [Paragraph("Должность", S["tbl_cell"]),
         Paragraph(data.position, S["tbl_cell"])],
        [Paragraph("Роль в системе", S["tbl_cell"]),
         Paragraph(data.role, S["tbl_cell"])],
        [Paragraph("Дата регистрации", S["tbl_cell"]),
         Paragraph(data.created_at.strftime("%d.%m.%Y"), S["tbl_cell"])],
    ]
    profile_t = Table(profile_data, colWidths=[6*cm, CW-6*cm])
    st = _tbl_base_style()
    for i in range(1, len(profile_data), 2):
        st.append(("BACKGROUND",(0,i),(-1,i),_WHITE))
    profile_t.setStyle(TableStyle(st))
    e.append(profile_t)

    # 2. Задания
    e.append(Paragraph("2. Выполнение заданий", S["section"]))
    summ_data = [
        [Paragraph("Показатель", S["tbl_hdr"]),
         Paragraph("Кол-во", S["tbl_hdr"]),
         Paragraph("Доля", S["tbl_hdr"])],
        [Paragraph("Всего назначено", S["tbl_cell"]),
         Paragraph(str(data.assignments_total), S["tbl_cell_c"]),
         Paragraph("100 %", S["tbl_cell_c"])],
        [Paragraph("Выполнено", S["tbl_cell"]),
         Paragraph(str(data.assignments_completed), S["tbl_cell_c"]),
         Paragraph(f"{data.completion_rate} %", S["tbl_cell_c"])],
        [Paragraph("В процессе", S["tbl_cell"]),
         Paragraph(str(data.assignments_in_progress), S["tbl_cell_c"]),
         Paragraph(f"{round(data.assignments_in_progress/max(data.assignments_total,1)*100,1)} %", S["tbl_cell_c"])],
        [Paragraph("Просрочено", S["tbl_cell"]),
         Paragraph(str(data.assignments_overdue), S["tbl_cell_c"]),
         Paragraph(f"{round(data.assignments_overdue/max(data.assignments_total,1)*100,1)} %", S["tbl_cell_c"])],
    ]
    summ_t = Table(summ_data, colWidths=[CW*0.5, CW*0.25, CW*0.25])
    st2 = _tbl_base_style()
    for i in range(1, len(summ_data), 2):
        st2.append(("BACKGROUND",(0,i),(-1,i),_WHITE))
    st2.append(("ALIGN",(1,0),(-1,-1),"CENTER"))
    summ_t.setStyle(TableStyle(st2))
    e.append(summ_t)

    # 3. AI-сессии
    score_str = f"{data.avg_dialog_score:.2f}" if data.avg_dialog_score else "—"
    e.append(Paragraph("3. Работа с AI-тренажёром", S["section"]))
    sess_data = [
        [Paragraph("Показатель", S["tbl_hdr"]),
         Paragraph("Значение", S["tbl_hdr"])],
        [Paragraph("Всего диалоговых сессий", S["tbl_cell"]),
         Paragraph(str(data.dialogs_total), S["tbl_cell_c"])],
        [Paragraph("Завершено успешно", S["tbl_cell"]),
         Paragraph(str(data.dialogs_completed), S["tbl_cell_c"])],
        [Paragraph("Средний балл (из 10)", S["tbl_cell"]),
         Paragraph(score_str, S["tbl_cell_c"])],
    ]
    sess_t = Table(sess_data, colWidths=[CW*0.65, CW*0.35])
    st3 = _tbl_base_style()
    for i in range(1, len(sess_data), 2):
        st3.append(("BACKGROUND",(0,i),(-1,i),_WHITE))
    st3.append(("ALIGN",(1,0),(-1,-1),"CENTER"))
    sess_t.setStyle(TableStyle(st3))
    e.append(sess_t)

    # 4. Навыки
    if data.skill_levels:
        e.append(Paragraph("4. Уровень освоения навыков", S["section"]))
        sk_rows = [[
            Paragraph("Навык", S["tbl_hdr"]),
            Paragraph("Категория", S["tbl_hdr"]),
            Paragraph("Уровень, %", S["tbl_hdr"]),
            Paragraph("Статус", S["tbl_hdr"]),
            Paragraph("Попыток", S["tbl_hdr"]),
        ]]
        for s in data.skill_levels:
            sk_rows.append([
                Paragraph(s["skill_name"], S["tbl_cell"]),
                Paragraph(s.get("category", "—"), S["tbl_cell"]),
                Paragraph(f"{s['current_level']:.0f}", S["tbl_cell_c"]),
                Paragraph(_mastery_label(s["mastery_status"]), S["tbl_cell_c"]),
                Paragraph(str(s["attempts_count"]), S["tbl_cell_c"]),
            ])
        sk_t = Table(sk_rows, colWidths=[CW*0.28, CW*0.24, CW*0.13, CW*0.22, CW*0.13], repeatRows=1)
        st4 = _tbl_base_style()
        for i in range(1, len(sk_rows), 2):
            st4.append(("BACKGROUND",(0,i),(-1,i),_WHITE))
        st4.append(("ALIGN",(2,0),(-1,-1),"CENTER"))
        sk_t.setStyle(TableStyle(st4))
        e.append(sk_t)

    # 5. Заключение
    e.append(Spacer(1, 14))
    e.append(HRFlowable(width="100%", thickness=0.5, color=_BLACK, spaceAfter=8))
    e.append(Paragraph("5. Заключение", S["section"]))
    conclusion = (
        f"По результатам "
        f"прохождения курса "
        f"сотрудник {data.short_name} "
        f"выполнил {data.assignments_completed} "
        f"из {data.assignments_total} назначенных "
        f"заданий ({data.completion_rate} %). "
        f"Проведено {data.dialogs_total} "
        f"диалоговых сессий, "
        f"из которых {data.dialogs_completed} "
        f"завершены успешно."
    )
    if data.avg_dialog_score:
        conclusion += (
            f" Средний балл "
            f"{data.avg_dialog_score:.2f} из 10."
        )
    e.append(Paragraph(conclusion, S["body"]))
    e.append(Spacer(1, 24))
    sig_data = [
        [Paragraph("Отчёт составил:", S["tbl_cell"]),
         Paragraph("______________ / " + data.generated_by + " /", S["tbl_cell"])],
        [Paragraph("Дата:", S["tbl_cell"]),
         Paragraph(data.generated_at.strftime("%d.%m.%Y"), S["tbl_cell"])],
    ]
    sig_t = Table(sig_data, colWidths=[4.5*cm, CW-4.5*cm])
    sig_t.setStyle(TableStyle([
        ("FONTNAME",(0,0),(-1,-1),"LiberSerif"),
        ("FONTSIZE",(0,0),(-1,-1),11),
        ("FONTNAME",(0,0),(0,-1),"LiberSerif-B"),
        ("LINEBELOW",(1,0),(1,0),0.5,_MGRAY),
        ("TOPPADDING",(0,0),(-1,-1),8),
        ("BOTTOMPADDING",(0,0),(-1,-1),8),
    ]))
    e.append(sig_t)
    e.append(Spacer(1, 16))
    e.append(HRFlowable(width="100%", thickness=0.5, color=_MGRAY, spaceAfter=4))
    e.append(Paragraph(
        f"Документ сформирован "
        f"автоматически • "
        f"ИС «ИИ-Академия» • "
        f"{data.generated_at.strftime('%d.%m.%Y %H:%M')}",
        S["footer"]))
    doc.build(e)
    return buf.getvalue()


# ════════════════════════════════════════════════════════════════════════
# DOCX
# ════════════════════════════════════════════════════════════════════════
def _doc_setup() -> Document:
    doc = Document()
    sect = doc.sections[0]
    sect.page_width    = DCm(21)
    sect.page_height   = DCm(29.7)
    sect.left_margin   = DCm(3)
    sect.right_margin  = DCm(1.5)
    sect.top_margin    = DCm(2)
    sect.bottom_margin = DCm(2)
    return doc


def _set_borders(cell, color="000000"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBd = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBd.append(el)
    tcPr.append(tcBd)


def _set_shade(cell, fill_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _para(cell, text: str, bold=False, center=False, size_pt=12):
    p = cell.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size_pt)
    r.font.name = "Times New Roman"
    r.font.color.rgb = RGBColor(0, 0, 0)


def _heading(doc: Document, text: str, numbered: str = "") -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.first_line_indent = DCm(0)
    r = p.add_run((numbered + " " + text).strip())
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = "Times New Roman"
    r.font.color.rgb = RGBColor(0, 0, 0)


def _body(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = DCm(1.25)
    r = p.add_run(text)
    r.font.size = Pt(14)
    r.font.name = "Times New Roman"
    r.font.color.rgb = RGBColor(0, 0, 0)


def _add_info_table(doc: Document, rows: list, col_widths_cm: list):
    tbl = doc.add_table(rows=0, cols=len(col_widths_cm))
    tbl.style = "Table Grid"
    for ri, row_data in enumerate(rows):
        row = tbl.add_row()
        for ci, txt in enumerate(row_data):
            cell = row.cells[ci]
            cell.width = DCm(col_widths_cm[ci])
            _set_borders(cell)
            if ri == 0:
                _set_shade(cell, "E8E8E8")
            elif ri % 2 == 0:
                _set_shade(cell, "F7F7F7")
            _para(cell, txt, bold=(ri == 0),
                  center=(ci > 0 and ri == 0) or (ci > 1),
                  size_pt=12)
    doc.add_paragraph()


def generate_docx(data: UserReportData) -> bytes:
    doc = _doc_setup()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)

    org_lines = [
        "Министерство "
        "науки и высшего "
        "образования "
        "Российской Федерации",
        "Федеральное "
        "государственное "
        "бюджетное образо"
        "вательное учреждение",
        "высшего образования",
        "«Новгородский "
        "государственный "
        "университет "
        "имени Ярослава "
        "Мудрого»",
    ]
    for line in org_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = DCm(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(line)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0, 0, 0)

    sep = doc.add_paragraph()
    sep.paragraph_format.first_line_indent = DCm(0)
    sep.paragraph_format.space_before = Pt(4)
    sep.paragraph_format.space_after  = Pt(4)
    pPr = sep._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    btm = OxmlElement("w:bottom")
    btm.set(qn("w:val"), "double")
    btm.set(qn("w:sz"), "6")
    btm.set(qn("w:space"), "1")
    btm.set(qn("w:color"), "000000")
    pBdr.append(btm)
    pPr.append(pBdr)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.first_line_indent = DCm(0)
    title_p.paragraph_format.space_before = Pt(16)
    title_p.paragraph_format.space_after  = Pt(4)
    r = title_p.add_run("ОТЧЁТ О ДОСТИЖЕНИЯХ")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0, 0, 0)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.first_line_indent = DCm(0)
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after  = Pt(16)
    r = sub_p.add_run(
        "при прохождении "
        "курса в системе "
        "«ИИ-Академия»"
    )
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0, 0, 0)

    _add_info_table(doc, [
        ["Реквизит", "Значение"],
        ["Составил", data.generated_by],
        ["Дата формирования",
         data.generated_at.strftime("%d.%m.%Y")],
    ], col_widths_cm=[6.5, 9.0])

    _heading(doc, "Сведения о сотруднике", "1.")
    _add_info_table(doc, [
        ["Показатель", "Значение"],
        ["Фамилия, имя, отчество", data.full_name],
        ["Электронная почта", data.email],
        ["Подразделение", data.department],
        ["Должность", data.position],
        ["Роль в системе", data.role],
        ["Дата регистрации",
         data.created_at.strftime("%d.%m.%Y")],
    ], col_widths_cm=[7.0, 8.5])

    _heading(doc, "Выполнение заданий", "2.")
    _add_info_table(doc, [
        ["Показатель", "Кол-во", "Доля, %"],
        ["Всего назначено",
         str(data.assignments_total), "100"],
        ["Выполнено",
         str(data.assignments_completed), str(data.completion_rate)],
        ["В процессе",
         str(data.assignments_in_progress),
         str(round(data.assignments_in_progress/max(data.assignments_total,1)*100,1))],
        ["Просрочено",
         str(data.assignments_overdue),
         str(round(data.assignments_overdue/max(data.assignments_total,1)*100,1))],
    ], col_widths_cm=[8.5, 2.5, 2.5])

    score_str = f"{data.avg_dialog_score:.2f}" if data.avg_dialog_score else "—"
    _heading(doc, "Работа с AI-тренажёром", "3.")
    _add_info_table(doc, [
        ["Показатель", "Значение"],
        ["Всего диалоговых сессий",
         str(data.dialogs_total)],
        ["Завершено успешно",
         str(data.dialogs_completed)],
        ["Средний балл (из 10)", score_str],
    ], col_widths_cm=[10.0, 3.5])

    if data.skill_levels:
        _heading(doc, "Уровень освоения навыков", "4.")
        sk_rows = [["Навык", "Категория",
                    "Уровень, %", "Статус", "Попыток"]]
        for s in data.skill_levels:
            sk_rows.append([
                s["skill_name"],
                s.get("category", "—"),
                f"{s['current_level']:.0f}",
                _mastery_label(s["mastery_status"]),
                str(s["attempts_count"]),
            ])
        _add_info_table(doc, sk_rows, col_widths_cm=[4.5, 3.5, 2.0, 3.0, 1.5])

    _heading(doc, "Заключение", "5.")
    conclusion = (
        f"По результатам "
        f"прохождения курса "
        f"сотрудник {data.short_name} "
        f"выполнил {data.assignments_completed} "
        f"из {data.assignments_total} назначенных "
        f"заданий ({data.completion_rate} %). "
        f"Проведено {data.dialogs_total} "
        f"диалоговых сессий, "
        f"из которых {data.dialogs_completed} "
        f"завершены успешно."
    )
    if data.avg_dialog_score:
        conclusion += f" Средний балл {data.avg_dialog_score:.2f} из 10."
    _body(doc, conclusion)

    doc.add_paragraph().paragraph_format.space_before = Pt(20)
    sign_tbl = doc.add_table(rows=2, cols=2)
    sign_tbl.style = "Table Grid"
    labels = ["Отчёт составил:", "Дата:"]
    vals   = ["______________ / " + data.generated_by + " /",
              data.generated_at.strftime("%d.%m.%Y")]
    for ri, (lbl, val) in enumerate(zip(labels, vals)):
        for ci, txt in enumerate([lbl, val]):
            cell = sign_tbl.cell(ri, ci)
            cell.width = DCm(7.5)
            _set_borders(cell, "CCCCCC")
            _para(cell, txt, bold=(ci == 0), size_pt=12)

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot.paragraph_format.first_line_indent = DCm(0)
    foot.paragraph_format.space_before = Pt(20)
    pPr2 = foot._p.get_or_add_pPr()
    pBdr2 = OxmlElement("w:pBdr")
    top2  = OxmlElement("w:top")
    top2.set(qn("w:val"), "single")
    top2.set(qn("w:sz"), "4")
    top2.set(qn("w:space"), "1")
    top2.set(qn("w:color"), "AAAAAA")
    pBdr2.append(top2)
    pPr2.append(pBdr2)
    r = foot.add_run(
        f"Документ сформирован "
        f"автоматически • "
        f"ИС «ИИ-Академия» • "
        f"{data.generated_at.strftime('%d.%m.%Y %H:%M')}"
    )
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(100, 100, 100)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
